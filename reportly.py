# reportly.py
from flask import request, render_template, redirect, url_for, flash, send_file, jsonify, Response
from werkzeug.utils import secure_filename
from pathlib import Path
from collections import defaultdict, deque
from docx import Document
from openai import OpenAI
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor
import uuid, json, os, requests, re, time, shutil
import subprocess
import tempfile, os
# --- Robust section & table parsing -----------------------------------------
import difflib
from docx.table import Table as DocxTable
from docx.oxml.ns import qn


# =========================
# OpenAI client (env)
# =========================
OPENAI_MODEL = "gpt-4"
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")

_client = None
def _get_openai_client():
    global _client
    if _client is None:
        _client = OpenAI(api_key=OPENAI_API_KEY, max_retries=1)
    return _client

# =========================
# Paths / storage
# =========================
UPLOAD_DIR = Path("uploads")
OUTPUT_DIR = Path("outputs")
STATE_DIR  = Path("state")
for d in (UPLOAD_DIR, OUTPUT_DIR, STATE_DIR):
    d.mkdir(parents=True, exist_ok=True)

GOTENBERG_URL = os.environ.get("GOTENBERG_URL")  # optional DOCX->PDF via Gotenberg

# =========================
# SSE (preview refresh)
# =========================
_EVENT_QUEUES = defaultdict(lambda: deque(maxlen=200))

def _push_event(doc_id: str, payload: dict):
    _EVENT_QUEUES[doc_id].append(json.dumps(payload))

def sse_stream(doc_id: str):
    """Server-Sent Events generator for /reportly/stream/<doc_id>"""
    last_ping = time.time()
    q = _EVENT_QUEUES[doc_id]
    idx = 0
    while True:
        if time.time() - last_ping > 15:
            yield "event: ping\ndata: {}\n\n"
            last_ping = time.time()
        if idx < len(q):
            data = q[idx]; idx += 1
            yield f"event: update\ndata: {data}\n\n"
        else:
            time.sleep(0.25)

# =========================
# Helpers
# =========================
def allowed_file(filename: str) -> bool:
    return Path(filename).suffix.lower() in {".docx", ".pdf"}

def _normalize_ws(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip()

def _key(s: str) -> str:
    """normalize heading title for fuzzy match."""
    return re.sub(r"[^a-z0-9]+", "", (s or "").lower())

def _para_style_name(p):
    try:
        return (p.style.name or "").lower()
    except Exception:
        return ""

def _is_heading(p):
    s = _para_style_name(p)
    return ("heading 1" in s) or ("heading 2" in s) or ("heading 3" in s)

def _heading_level(p):
    s = _para_style_name(p)
    if "heading 1" in s: return 1
    if "heading 2" in s: return 2
    if "heading 3" in s: return 3
    return 99

def _insert_par_after(anchor: Paragraph) -> Paragraph:
    """Insert an empty paragraph after `anchor`."""
    new_p = OxmlElement("w:p")
    anchor._element.addnext(new_p)
    return Paragraph(new_p, anchor._parent)

def _list_headings(doc: Document):
    """Return list of (idx, level, text, key) in doc order."""
    out = []
    for i, p in enumerate(doc.paragraphs):
        if _is_heading(p):
            txt = _normalize_ws(p.text)
            out.append((i, _heading_level(p), txt, _key(txt)))
    return out

HEADING_STYLE_PREFIXES = ("heading", "encabezado", "überschrift", "rubrique", "titre", "título")

def _is_heading_style(p):
    s = _para_style_name(p)
    if not s:
        return False
    s = s.lower()
    return any(s.startswith(hs) for hs in HEADING_STYLE_PREFIXES) or re.search(r"\bheading\s*[1-6]\b", s)

def _infer_heading_level(p):
    s = _para_style_name(p)
    for lvl in range(1, 7):
        if f"heading {lvl}" in s:
            return lvl
    txt = _normalize_ws(p.text)
    if re.match(r"^\d+(\.\d+)*\s+\S", txt) or txt.endswith(":") or (txt.isupper() and 3 <= len(txt) <= 80):
        return 2
    return 99

def _iter_block_items(document):
    parent_elm = document._element.body
    for child in parent_elm.iterchildren():
        tag = child.tag.rsplit('}', 1)[-1]
        if tag == "p":
            yield Paragraph(child, document)
        elif tag == "tbl":
            yield DocxTable(child, document)

def _table_to_json(tbl: DocxTable):
    headers, rows = [], []
    if tbl.rows:
        first = [c.text.strip() for c in tbl.rows[0].cells]
        looks_header = any(first) and all(len(h) <= 80 for h in first)
        if looks_header:
            headers = first
            data_rows = tbl.rows[1:]
        else:
            data_rows = tbl.rows
        for r in data_rows:
            rows.append([c.text.strip() for c in r.cells])
    return {"type": "table", "headers": headers, "rows": rows}

def _para_to_block(p: Paragraph):
    t = _normalize_ws(p.text)
    if not t:
        return None
    style = _para_style_name(p)
    if "list" in style:
        mt = re.match(r"^(?:[\u2022\-\*]|\d+\.)\s+(.*)$", t)
        core = mt.group(1) if mt else t
        kind = "number" if re.match(r"^\d+\.", t) else "bullet"
        return {"type": "list", "style": kind, "items": [core]}
    return {"type": "paragraph", "text": t}

def parse_docx_sections_full(doc: Document):
    sections, current = [], None
    def ensure_current(title="Preamble", level=1):
        nonlocal current
        if current is None:
            current = {"id": str(uuid.uuid4()), "title": title, "level": level, "style_name": "Heading 1", "blocks": []}
            sections.append(current)

    for item in _iter_block_items(doc):
        if isinstance(item, Paragraph):
            if _is_heading_style(item):
                lvl = _infer_heading_level(item)
                title = _normalize_ws(item.text) or f"Section {len(sections)+1}"
                current = {"id": str(uuid.uuid4()), "title": title, "level": min(lvl, 6), "style_name": item.style.name if item.style else "", "blocks": []}
                sections.append(current)
            else:
                ensure_current()
                blk = _para_to_block(item)
                if blk:
                    if blk["type"] == "list" and current["blocks"] and current["blocks"][-1]["type"] == "list" and current["blocks"][-1]["style"] == blk["style"]:
                        current["blocks"][-1]["items"].extend(blk["items"])
                    else:
                        current["blocks"].append(blk)
        else:  # table
            ensure_current()
            current["blocks"].append(_table_to_json(item))

    if not sections:
        sections.append({"id": str(uuid.uuid4()), "title": "Document", "level": 1, "style_name": "Heading 1", "blocks": []})
    return sections

# =========================
# Compose (from scratch)
# =========================
def compose_from_structure(template_path: Path, structure: list, outfile: Path, strip_existing=True):
    doc = Document(str(template_path))
    if strip_existing:
        clear_document_body(doc)

    for si, sec in enumerate(structure):
        title = sec.get("title", "")
        if title:
            ph = doc.add_paragraph(title)
            level = int(sec.get("level", 1))
            style_name = "Heading 1" if level == 1 else "Heading 2" if level == 2 else "Heading 3"
            try: ph.style = doc.styles[style_name]
            except KeyError: pass

        for block in sec.get("blocks", []):
            btype = block.get("type")
            if btype == "paragraph":
                text = _normalize_ws(block.get("text", ""))
                if text:
                    p = doc.add_paragraph(text)
                    try: p.style = doc.styles["Normal"]
                    except KeyError: pass
            elif btype == "list":
                items = block.get("items", []); style = block.get("style", "bullet")
                for idx, it in enumerate(items):
                    txt = _normalize_ws(it)
                    if not txt: continue
                    p = doc.add_paragraph(f"{idx+1}. {txt}" if style == "number" else txt)
                    try: p.style = doc.styles["List Paragraph"]
                    except KeyError: pass
            elif btype == "table":
                headers = block.get("headers", []); rows = block.get("rows", [])
                cols = len(headers) if headers else (len(rows[0]) if rows else 0)
                if cols > 0:
                    t = doc.add_table(rows=1 if headers else 0, cols=cols)
                    if headers:
                        for i, h in enumerate(headers[:cols]): t.rows[0].cells[i].text = _normalize_ws(h)
                    for r in rows:
                        row_cells = t.add_row().cells
                        for i, cell_val in enumerate(r[:cols]): row_cells[i].text = _normalize_ws(cell_val)

        if si < len(structure) - 1:
            doc.add_paragraph("")
    doc.save(str(outfile))

# =========================
# Edit-in-place (ordinal-aware)
# =========================
def _find_anchor_index(doc: Document, heading_text: str, level: int, approx_order_index):
    heads = _list_headings(doc)
    if not heads:
        return -1, 99

    key = _key(heading_text or "")
    # 1) exact key match
    for idx, lvl, txt, kk in heads:
        if kk == key:
            return idx, lvl

    # 2) fuzzy title match
    best = (-1, 0.0, 99)
    for idx, lvl, txt, kk in heads:
        score = difflib.SequenceMatcher(None, txt.lower(), (heading_text or "").lower()).ratio()
        if score > best[1]:
            best = (idx, score, lvl)
    if best[0] != -1 and best[1] >= 0.72:
        return best[0], best[2]

    # 3) same-level ordinal
    if approx_order_index is not None:
        same = [h for h in heads if h[1] == level]
        if same:
            j = min(approx_order_index, len(same)-1)
            return same[j][0], same[j][1]

    return -1, 99

def _insert_missing_heading(doc: Document, title: str, level: int) -> Paragraph:
    anchor_ref = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()
    new_p = OxmlElement("w:p")
    anchor_ref._element.addnext(new_p)
    hp = Paragraph(new_p, anchor_ref._parent)
    hp.text = title or "New Section"
    try:
        hp.style = doc.styles["Heading 1" if level==1 else "Heading 2" if level==2 else "Heading 3"]
    except Exception:
        pass
    return hp

def replace_section_content_in_place(doc: Document, heading_text: str, blocks: list, approx_order_index=None, level_hint: int | None = None):
    level_for_search = level_hint if level_hint else 1
    anchor_idx, _ = _find_anchor_index(doc, heading_text, level=level_for_search, approx_order_index=approx_order_index)

    if anchor_idx == -1:
        anchor = _insert_missing_heading(doc, heading_text, level_for_search)
        # Find the new heading by text, fallback to last paragraph
        anchor_idx = next(
            (i for i, p in enumerate(doc.paragraphs) if _normalize_ws(p.text) == _normalize_ws(heading_text)),
            len(doc.paragraphs) - 1
        )
        anchor = doc.paragraphs[anchor_idx]
    else:
        anchor = doc.paragraphs[anchor_idx]
        body = anchor._p.getparent()
        el = anchor._p.getnext()
        while el is not None:
            tag = el.tag.rsplit("}", 1)[-1]
            if tag == "p":
                temp_p = Paragraph(el, anchor._parent)
                if _is_heading(temp_p) and _heading_level(temp_p) <= _heading_level(anchor):
                    break
                nxt = el.getnext(); body.remove(el); el = nxt; continue
            elif tag == "tbl":
                nxt = el.getnext(); body.remove(el); el = nxt; continue
            el = el.getnext()

    insert_after = anchor
    for block in blocks:
        btype = block.get("type")
        if btype == "paragraph":
            p = _insert_par_after(insert_after); p.text = _normalize_ws(block.get("text",""))
            try:
                p.style = doc.styles["Normal"]
            except KeyError:
                pass
            insert_after = p
        elif btype == "list":
            items = block.get("items", []); style = block.get("style", "bullet")
            for idx, it in enumerate(items):
                p = _insert_par_after(insert_after)
                p.text = f"{idx+1}. {_normalize_ws(it)}" if style == "number" else _normalize_ws(it)
                try:
                    p.style = doc.styles["List Paragraph"]
                except KeyError:
                    pass
                insert_after = p
        elif btype == "table":
            headers = block.get("headers", []); rows = block.get("rows", [])
            cols = len(headers) if headers else (len(rows[0]) if rows else 0)
            if cols > 0:
                t = doc.add_table(rows=1 if headers else 0, cols=cols)
                if headers:
                    for i, h in enumerate(headers[:cols]): t.rows[0].cells[i].text = _normalize_ws(h)
                for r in rows:
                    row_cells = t.add_row().cells
                    for i, cell_val in enumerate(r[:cols]): row_cells[i].text = _normalize_ws(cell_val)
                tbl_el = t._tbl; body = insert_after._p.getparent(); body.remove(tbl_el); insert_after._element.addnext(tbl_el)
                insert_after = _insert_par_after(insert_after)

    return anchor_idx

def _apply_blocks_with_marker(doc: Document, title: str, blocks: list, approx_order_index, level_hint: int | None):
    idx = replace_section_content_in_place(doc, title, blocks, approx_order_index=approx_order_index, level_hint=level_hint)
    if idx < 0: return
    anchor = doc.paragraphs[idx]
    el = anchor._p.getnext()
    while el is not None:
        tag = el.tag.rsplit("}", 1)[-1]
        if tag == "p":
            temp = Paragraph(el, anchor._parent)
            if _is_heading(temp): break
            for run in temp.runs:
                try: run.font.color.rgb = RGBColor(0x0A, 0xA3, 0x68)
                except Exception: pass
            el = el.getnext(); continue
        if tag == "tbl":
            el = el.getnext(); continue
        el = el.getnext()

def build_review_preview_pdf(doc_id: str) -> Path:
    """
    Build preview with pending changes marked in green.
    """
    state = load_state(doc_id)
    if not state:
        raise FileNotFoundError("No state for preview")

    base_path = Path(state["working_path"])
    if not base_path.exists():
        raise FileNotFoundError("Working DOCX missing")

    # copy working to preview-docx, apply pending with green marks
    preview_docx = OUTPUT_DIR / f"{doc_id}_preview.docx"
    shutil.copyfile(base_path, preview_docx)

    pending = state.get("pending") or []
    if pending:
        d = Document(str(preview_docx))
        for change in pending:
            _apply_blocks_with_marker(
                d,
                change.get("title",""),
                change.get("blocks",[]),
                approx_order_index=change.get("order_index"),
                level_hint=change.get("level",1)
            )
        d.save(str(preview_docx))

    preview_pdf = OUTPUT_DIR / f"{doc_id}_preview.pdf"
    docx_to_pdf(preview_docx, preview_pdf)
    _push_event(doc_id, {"preview": "ready", "ts": int(time.time()*1000)})
    return preview_pdf

def build_preview_pdf(doc_id: str) -> Path:
    # route callers through review preview
    return build_review_preview_pdf(doc_id)

# =========================
# Conversions
# =========================
# reportly.py
def _pdf_to_docx_local(pdf_path: Path, out_docx: Path):
    import subprocess, time
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    before = {p.name for p in out_docx.parent.glob("*.docx")}

    cmd = [
        "soffice","--headless","--convert-to","docx",
        "--outdir", str(out_docx.parent), str(pdf_path)
    ]
    proc = subprocess.run(cmd, capture_output=True, check=False)
    if proc.returncode != 0:
        raise RuntimeError(proc.stderr.decode('utf-8','ignore') or proc.stdout.decode('utf-8','ignore'))

    time.sleep(0.2)  # FS settle
    after = [p for p in out_docx.parent.glob("*.docx") if p.name not in before]
    if not after:
        raise FileNotFoundError("Conversion succeeded but no DOCX appeared.")
    # prefer same-stem, else newest
    cand = sorted(after, key=lambda p: p.stat().st_mtime, reverse=True)
    target = next((p for p in cand if p.stem.startswith(Path(pdf_path).stem)), cand[0])
    if out_docx.exists(): out_docx.unlink()
    target.rename(out_docx)
    return out_docx

def _docx_to_pdf_local(docx_path: Path, out_pdf: Path):
    out_pdf.parent.mkdir(parents=True, exist_ok=True)
    cmd = [
        "soffice","--headless","--convert-to","pdf",
        "--outdir", str(out_pdf.parent), str(docx_path)
    ]
    proc = subprocess.run(cmd, capture_output=True, check=False)
    if proc.returncode != 0:
        raise RuntimeError(proc.stderr.decode('utf-8','ignore') or proc.stdout.decode('utf-8','ignore'))
    produced = docx_path.with_suffix(".pdf")
    produced_path = out_pdf.parent / produced.name
    if produced_path != out_pdf:
        if out_pdf.exists(): out_pdf.unlink()
        produced_path.rename(out_pdf)
    return out_pdf

def docx_to_pdf(docx_path: Path, out_pdf: Path):
    if GOTENBERG_URL:
        with open(docx_path, "rb") as f:
            files = {"files": ("document.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            r = requests.post(f"{GOTENBERG_URL}/forms/libreoffice/convert", files=files, timeout=120)
            r.raise_for_status()
            out_pdf.parent.mkdir(parents=True, exist_ok=True)
            out_pdf.write_bytes(r.content)
            return out_pdf
    return _docx_to_pdf_local(docx_path, out_pdf)

# =========================
# LLM prompting
# =========================
LLM_SYSTEM_PROMPT = """You generate regulated enterprise report content.

Return ONLY a single JSON object that matches this schema EXACTLY. Do not wrap your answer in code fences. Do not add any commentary, prefixes, or suffixes.

{
  "sections": [
    {
      "id": "string (echo the provided id)",
      "title": "string (echo the provided title unless instructed to rename)",
      "level": 1|2|3,
      "blocks": [
        { "type": "paragraph", "text": "plain text, no markdown" } |
        { "type": "list", "style": "bullet"|"number", "items": ["..."] } |
        { "type": "table", "headers": ["..."], "rows": [["..."]]}
      ]
    }
  ]
}

Rules:
- Respect the given section order and levels.
- Do NOT include markdown, HTML, code fences, or style tags in text.
- Stay within the requested length mode (concise/standard/detailed/custom words).
- Keep tables <= 8 columns. Avoid empty blocks.
- Output MUST be valid JSON and parse with json.loads()."""

def build_user_llm_message(prompt_text, structure, include_source_text=False, source_text=None, length_mode="standard", custom_words=None):
    skeleton = [{"id": s["id"], "title": s["title"], "level": s["level"]} for s in structure]
    payload = {
        "instruction": prompt_text,
        "length_mode": length_mode,
        "custom_word_count": custom_words,
        "template_structure": skeleton,
        "include_source_text": include_source_text,
        "source_text_excerpt": source_text[:6000] if (include_source_text and source_text) else None
    }
    msg = {"role": "user", "content": json.dumps(payload, ensure_ascii=False)}
    msg["template_structure"] = skeleton  # for offline fallback
    return msg

def extract_source_plaintext(doc: Document, max_chars=15000):
    texts, size = [], 0
    for p in doc.paragraphs:
        t = _normalize_ws(p.text)
        if t:
            texts.append(t)
            size += len(t) + 1
            if size > max_chars: break
    return "\n".join(texts)[:max_chars]

def _repair_json_loose(s: str) -> str:
    s = s.strip()
    if s.startswith("```"):
        s = re.sub(r"^```(?:json)?\s*", "", s)
        s = re.sub(r"\s*```$", "", s)
    s = s.replace("\ufeff", "")
    s = re.sub(r",\s*([\]}])", r"\1", s)
    return s

def call_openai_structured(messages, timeout=120):
    # helper: build a realistic fallback using the provided template
    def _fallback():
        structure, instruction, length_mode = _extract_template_and_meta(messages)
        return _generate_demo_sections(structure, instruction, length_mode)

    # If there is no key at all, still give usable staged edits
    if not OPENAI_API_KEY:
        return _fallback()

    # Try the real call; on ANY null/empty/invalid result, auto-fallback
    try:
        client = _get_openai_client()
        safe_messages = []
        for m in messages:
            role = m.get("role", "user")
            content = m.get("content", "")
            if not isinstance(content, str):
                content = json.dumps(content, ensure_ascii=False)
            safe_messages.append({"role": role, "content": content})

        resp = client.chat.completions.create(
            model=OPENAI_MODEL, messages=safe_messages,
            temperature=0.2, max_tokens=2000, timeout=timeout
        )

        content = None
        if resp and getattr(resp, "choices", None):
            msg = resp.choices[0].message
            content = getattr(msg, "content", None)

        # Null, tool-only, or empty payload → fallback
        if _is_nullish_content(content):
            return _fallback()

        # Parse JSON; if busted or empty → fallback
        try:
            obj = json.loads(content)
        except Exception:
            try:
                obj = json.loads(_repair_json_loose(content))
            except Exception:
                m = re.search(r"\{[\s\S]*\}", content or "")
                if not m:
                    return _fallback()
                try:
                    obj = json.loads(_repair_json_loose(m.group(0)))
                except Exception:
                    return _fallback()

        if not obj or not isinstance(obj, dict) or not obj.get("sections"):
            return _fallback()

        return obj

    except Exception:
        # Network errors, API errors, etc. → fallback so UX keeps flowing
        return _fallback()

def pick_target_sections(user_text: str, structure: list):
    text = (user_text or "").lower()
    chosen = []
    for sec in structure:
        title = (sec.get("title") or "").lower()
        if not title: continue
        tokens = [t for t in re.split(r"[^a-z0-9]+", title) if len(t) >= 4]
        if any(tok in text for tok in tokens):
            chosen.append(sec)
    return chosen if chosen else structure

# =========================
# State helpers
# =========================

def _doc_to_fast_html(doc):
    """Basic HTML outline generator for docx Document."""
    html = []
    for p in doc.paragraphs:
        txt = _normalize_ws(p.text)
        if not txt:
            continue
        style = _para_style_name(p)
        if _is_heading(p):
            lvl = _heading_level(p)
            html.append(f'<h{lvl}>{txt}</h{lvl}>')
        else:
            html.append(f'<p>{txt}</p>')
    return "\n".join(html)

def save_state(doc_id: str, data: dict):
    path = STATE_DIR / f"{doc_id}.json"
    tmp  = path.with_suffix(".json.tmp")
    tmp.write_text(json.dumps(data, indent=2))
    os.replace(tmp, path)  # ✅ atomic
    # REMOVE any extra write here


def load_state(doc_id: str):
    f = STATE_DIR / f"{doc_id}.json"
    return json.loads(f.read_text()) if f.exists() else None

# =========================
# Routes (used by app.py)
# =========================
def handle_upload():
    file = request.files.get("file")
    if not file or file.filename == "":
        flash("Please choose a .docx or .pdf file to upload.", "error")
        return render_template("reportly.html", state=None)
    if not allowed_file(file.filename):
        flash("Only .docx or .pdf files are supported.", "error")
        return render_template("reportly.html", state=None)

    doc_id = str(uuid.uuid4())
    original_name = secure_filename(file.filename)
    up_path = UPLOAD_DIR / f"{doc_id}_{original_name}"
    file.save(up_path)

    # PDF → DOCX
    if up_path.suffix.lower() == ".pdf":
        converted = UPLOAD_DIR / f"{doc_id}_converted.docx"
        try:
            _pdf_to_docx_local(up_path, converted)
        except Exception as e:
            flash(f"PDF conversion failed: {e}", "error")
            return render_template("reportly.html", state=None)
        up_path = converted
        original_name = Path(original_name).stem + ".docx"
    # parse
    try:
        doc = Document(str(up_path))
    except Exception as e:
        flash(f"Failed to parse .docx: {e}", "error")
        return render_template("reportly.html", state=None)

    working_path = OUTPUT_DIR / f"{doc_id}_working.docx"
    shutil.copyfile(up_path, working_path)

    structure  = parse_docx_sections_full(doc)
    fast_html  = _doc_to_fast_html(doc)

    state = {
        "doc_id": doc_id,
        "filename": original_name,
        "basename": Path(original_name).stem,
        "upload_path": str(up_path),
        "working_path": str(working_path),
        "strip_existing": False,
        "include_in_context": False,
        "structure": structure,
        "messages": [],
        "html_outline": fast_html,
        "pending": []
    }
    save_state(doc_id, state)

    # initial preview (just the working doc)
    try:
        build_preview_pdf(doc_id)
    except Exception as e:
        flash(f"Preview build warning: {e}", "error")

    return redirect(url_for("reportly_edit", doc_id=doc_id))

def handle_edit(doc_id):
    state = load_state(doc_id)
    if not state:
        flash("Document not found.", "error")
        return redirect(url_for("reportly_home"))
    return render_template("reportly.html", state=state)

_LAST_BUILD = {}
def _should_build(doc_id, ms=300):
    t = int(time.time()*1000)
    last = _LAST_BUILD.get(doc_id, 0)
    if t - last < ms: return False
    _LAST_BUILD[doc_id] = t
    return True

def handle_state_update(doc_id):
    state = load_state(doc_id)
    if not state:
        return jsonify({"ok": False, "error": "Document not found"}), 404
    if "include_in_context" in request.form:
        state["include_in_context"] = request.form.get("include_in_context") in ("1","true","on")
    if "strip_existing" in request.form:
        state["strip_existing"] = request.form.get("strip_existing") in ("1","true","on")
    save_state(doc_id, state)
    try:
        build_preview_pdf(doc_id)
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
    if _should_build(doc_id): build_preview_pdf(doc_id)
    return jsonify({"ok": True})

def handle_download_current(doc_id):
    state = load_state(doc_id)
    if not state:
        flash("Document not found.", "error")
        return redirect(url_for("reportly_home"))
    return send_file(
        state["working_path"],
        as_attachment=True,
        download_name=state["filename"],
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def handle_preview(doc_id):
    pdf_path = OUTPUT_DIR / f"{doc_id}_preview.pdf"
    if not pdf_path.exists():
        try:
            pdf_path = build_preview_pdf(doc_id)
        except Exception as e:
            flash(f"Preview failed: {e}", "error")
            state = load_state(doc_id)
            return render_template("reportly.html", state=state or None)
    if request.args.get("download") == "1":
        return send_file(str(pdf_path), as_attachment=True, download_name=f"{load_state(doc_id)['basename']}.pdf")
    resp = send_file(str(pdf_path), mimetype="application/pdf")
    resp.headers["Cache-Control"] = "no-store, max-age=0"
    return resp

def handle_chat(doc_id):
    state = load_state(doc_id)
    if not state:
        return jsonify({"ok": False, "error": "Document not found"}), 404

    raw_message = (request.form.get("message") or "").strip()
    if not raw_message:
        return jsonify({"ok": False, "error": "Message is required"}), 400

    length_mode  = (request.form.get("length_mode") or "standard").lower()
    custom_words = request.form.get("custom_words")
    try:
        cw = int(custom_words) if (custom_words and custom_words.isdigit()) else None
    except Exception:
        cw = None

    # Optional: expand [print:<uuid>] token into context (saved via /api/print/upload)
    bundle_text = None
    message_for_prompt = raw_message
    try:
        m = re.search(r"\[print:([0-9a-fA-F-]{36})\]", raw_message)
        if m:
            uid = m.group(1)
            bundle_path = OUTPUT_DIR / f"{uid}_print.txt"
            if bundle_path.exists():
                bundle_text = bundle_path.read_text(encoding="utf-8", errors="ignore")[:120_000]
            message_for_prompt = re.sub(r"\[print:[^\]]+\]", "", raw_message).strip()
    except Exception:
        pass

    # LLM message history
    llm_messages = [{"role": "system", "content": LLM_SYSTEM_PROMPT}]
    for m in state.get("messages", []):
        llm_messages.append({"role": m["role"], "content": m["content"]})

    include_ctx = bool(state.get("include_in_context"))
    source_excerpt = None
    try:
        if include_ctx and state.get("working_path"):
            source_excerpt = extract_source_plaintext(Document(state["working_path"]))
    except Exception:
        source_excerpt = None

    full_structure   = state["structure"]
    target_structure = pick_target_sections(message_for_prompt, full_structure)

    # capture ordinal per level for better anchoring later
    level_counts = {1:0,2:0,3:0}
    order_map = {}
    for s in full_structure:
        lvl = int(s["level"])
        order_map[s["id"]] = level_counts[lvl]
        level_counts[lvl] += 1

    llm_user_msg = build_user_llm_message(
        prompt_text=message_for_prompt,
        structure=target_structure,
        include_source_text=include_ctx,
        source_text=source_excerpt,
        length_mode=length_mode,
        custom_words=cw
    )

    if bundle_text:
        try:
            payload = json.loads(llm_user_msg["content"])
            payload["external_materials"] = {
                "print_bundle_text": bundle_text,
                "note": "Materials collected via .print (tree + text snippets). Use only if relevant."
            }
            llm_user_msg["content"] = json.dumps(payload, ensure_ascii=False)
        except Exception:
            pass

    # persist user message
    state["messages"].append({"role": "user", "content": raw_message})
    save_state(doc_id, state)

    try:
        llm_messages.append(llm_user_msg)
        result = call_openai_structured(llm_messages)

        # Stage (not apply) as pending
        new_sections = result.get("sections", []) or []
        _ensure_pending(state)
        staged = []
        for s in new_sections:
            staged.append({
                "change_id": str(uuid.uuid4()),
                "title": s.get("title",""),
                "level": int(s.get("level",1)),
                "blocks": s.get("blocks", []),
                "order_index": order_map.get(s.get("id"), None)
            })
        state["pending"].extend(staged)
        save_state(doc_id, state)

        # Build preview showing green highlights
        build_preview_pdf(doc_id)

        assistant_msg = f"Proposed {len(staged)} change(s). Review on the right → Accept / Reject."
        state["messages"].append({"role": "assistant", "content": assistant_msg})
        save_state(doc_id, state)

        return jsonify({"ok": True, "assistant_message": assistant_msg})

    except Exception as e:
        err = f"Generation failed: {e}"
        state["messages"].append({"role": "assistant", "content": err})
        save_state(doc_id, state)
        return jsonify({"ok": False, "error": err}), 500

def handle_compose(doc_id):
    """Kept for compatibility; composes from structure into a fresh DOCX."""
    state = load_state(doc_id)
    if not state:
        flash("Document not found.", "error")
        return redirect(url_for("reportly_home"))

    out_path = OUTPUT_DIR / f"{doc_id}_composed.docx"
    if state.get("strip_existing", True):
        compose_from_structure(
            template_path=Path(state["upload_path"]),
            structure=state["structure"],
            outfile=out_path,
            strip_existing=True
        )
    else:
        compose_from_structure(
            template_path=Path(state["upload_path"]),
            structure=state["structure"],
            outfile=out_path,
            strip_existing=False
        )

    state["composed_path"] = str(out_path)
    save_state(doc_id, state)
    flash("Composed document created.", "success")
    return redirect(url_for("reportly_download", doc_id=doc_id))

def handle_download(doc_id):
    state = load_state(doc_id)
    if not state or not state.get("composed_path"):
        flash("No composed document available. Click 'Compose & Download' first.", "error")
        return redirect(url_for("reportly_edit", doc_id=doc_id))
    return send_file(state["composed_path"], as_attachment=True, download_name="report.docx")

# -------- REVIEW API --------
def handle_pending_list(doc_id):
    state = load_state(doc_id) or {}
    return jsonify({"ok": True, "pending": state.get("pending", [])})

def handle_review_accept(doc_id):
    state = load_state(doc_id)
    if not state: return jsonify({"ok": False, "error":"Document not found"}), 404
    _ensure_pending(state)
    body = request.get_json(silent=True) or {}
    change_id = body.get("change_id")
    apply_all = bool(body.get("all"))

    if apply_all:
        to_apply = list(state["pending"])
        state["pending"].clear()
    else:
        to_apply = [c for c in state["pending"] if c["change_id"] == change_id]
        state["pending"] = [c for c in state["pending"] if c["change_id"] != change_id]

    working = Path(state["working_path"])
    if not working.exists():
        return jsonify({"ok": False, "error": "Working DOCX missing"}), 500

    d = Document(str(working))
    for c in to_apply:
        replace_section_content_in_place(
            d,
            c.get("title",""),
            c.get("blocks", []),
            approx_order_index=c.get("order_index"),
            level_hint=c.get("level", 1)
        )
    d.save(str(working))
    save_state(doc_id, state)
    build_preview_pdf(doc_id)
    return jsonify({"ok": True, "applied": len(to_apply)})

def handle_review_reject(doc_id):
    state = load_state(doc_id)
    if not state: return jsonify({"ok": False, "error":"Document not found"}), 404
    _ensure_pending(state)
    body = request.get_json(silent=True) or {}
    change_id = body.get("change_id")
    reject_all = bool(body.get("all"))
    if reject_all:
        n = len(state["pending"])
        state["pending"].clear()
    else:
        before = len(state["pending"])
        state["pending"] = [c for c in state["pending"] if c["change_id"] != change_id]
        n = before - len(state["pending"])
    save_state(doc_id, state)
    build_preview_pdf(doc_id)
    return jsonify({"ok": True, "rejected": n})

def handle_structure_get(doc_id):
    state = load_state(doc_id)
    if not state:
        return jsonify({"ok": False, "error": "not found"}), 404
    items = []
    for s in state.get("structure", []):
        tables = sum(1 for b in s.get("blocks", []) if b.get("type") == "table")
        paras  = sum(1 for b in s.get("blocks", []) if b.get("type") in ("paragraph","list"))
        items.append({"id": s["id"], "title": s["title"], "level": int(s["level"]), "paragraphs": paras, "tables": tables})
    return jsonify({"ok": True, "sections": items})

# --- Null-safe fallback generators ------------------------------------------

def _ensure_pending(state):
    """Ensure the 'pending' key exists in the state dictionary."""
    if "pending" not in state or not isinstance(state["pending"], list):
        state["pending"] = []

def _demo_paragraph(title: str, instruction: str, length_mode: str) -> str:
    base = (f"{title} — demo content inserted automatically because the AI "
            f"response was empty. This lets you test Review → Accept and downloads.")
    if instruction:
        base += f" Instruction hint: {instruction[:160]}."
    if length_mode == "detailed":
        return base + " Includes extra detail to validate formatting and spacing."
    if length_mode == "concise":
        return base
    return base + " Key bullets follow for verification."

def _demo_bullets(title: str) -> dict:
    root = title.strip() or "Section"
    return {"type":"list","style":"bullet","items":[
        f"{root}: objective stated",
        f"{root}: assumptions captured",
        f"{root}: next actions with owners"
    ]}

def _demo_table(title: str) -> dict:
    return {"type":"table","headers":["Item","Owner","ETA"],"rows":[
        ["Kickoff","PM","Week 1"],
        ["Draft v1","Writer","Week 2"],
        ["Review","Leads","Week 3"],
    ]}

def _needs_table(title: str) -> bool:
    t = (title or "").lower()
    return any(k in t for k in ("timeline","schedule","plan","milestone","phases","budget","risk"))

def _generate_demo_sections(structure: list, instruction: str, length_mode: str) -> dict:
    out = []
    for i, s in enumerate(structure, 1):
        title = s.get("title") or f"Section {i}"
        level = int(s.get("level", 1))
        blocks = [
            {"type":"paragraph","text": _demo_paragraph(title, instruction, (length_mode or "standard").lower())},
            _demo_bullets(title),
        ]
        if _needs_table(title):
            blocks.append(_demo_table(title))
        out.append({"id": s["id"], "title": title, "level": level, "blocks": blocks})
    return {"sections": out}

def _extract_template_and_meta(messages):
    """Pull template_structure + meta from the last user message, even if it's a JSON string."""
    try:
        last = messages[-1]
        payload = last.get("content", "")
        if not isinstance(payload, str):
            payload = json.dumps(payload, ensure_ascii=False)
        data = json.loads(payload)
    except Exception:
        data = {}
    structure   = data.get("template_structure") or messages[-1].get("template_structure") or []
    instruction = data.get("instruction","")
    length_mode = (data.get("length_mode") or "standard").lower()
    return structure, instruction, length_mode

def _is_nullish_content(content: str | None) -> bool:
    if content is None:
        return True
    t = str(content).strip().lower()
    return t in ("", "null", "none", "{}", "[]", '"null"', '""')
