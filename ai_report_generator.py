#!/usr/bin/env python3
"""
ai_report_generator.py

AI-powered capital planning report generator using OpenAI API
"""

import os
import json
import datetime
from typing import Dict, List, Optional
from openai import OpenAI
import logging
import traceback
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Set OpenAI API key globally

# Set up logger
logger = logging.getLogger("vibeops.ai_report_generator")
logger.setLevel(logging.DEBUG)
if not logger.hasHandlers():
    handler = logging.StreamHandler()
    formatter = logging.Formatter('[%(asctime)s] %(levelname)s %(name)s: %(message)s')
    handler.setFormatter(formatter)
    logger.addHandler(handler)

# from openai import OpenAI
# client = OpenAI(api_key=os.getenv('OpenAI_API_KEY'))

# Restore and update AIReportGenerator
class AIReportGenerator:
    def __init__(self):
        logger.debug("Initializing AIReportGenerator...")
        self.api_key = os.getenv('OpenAI_API_KEY')
        if self.api_key:
            try:
                self.client = OpenAI(api_key=self.api_key)
                self.use_new_client = True
                logger.debug("Using OpenAI client.")
            except Exception as e:
                logger.warning(f"Could not initialize OpenAI client: {e}")
                self.client = None
                self.use_new_client = False
        else:
            logger.warning("No OpenAI API key found. AI features will be disabled.")
            self.client = None
            self.use_new_client = False

    def generate_report_sections(self, report_type: str, company_name: str, project_context: str = "") -> dict:
        logger.debug(f"Generating AI report sections for {report_type}, {company_name}, {project_context}")
        try:
            template = REPORT_TEMPLATES[report_type]
            table_headers = template['table_headers']
            # Build prompt for AI to generate all sections and table data
            prompt = f"""
You are a VibeOps {report_type.replace('_', ' ')} expert. Generate a detailed report for {company_name} with project context: {project_context}.
Return a JSON object with:
- executive_summary: string
- recommendations: list of strings
- {'risk_assessment: string' if report_type in ['capital_planning', 'feasibility_study'] else 'project_review: string'}
- {'financial_analysis: string' if report_type == 'capital_planning' else 'findings: string' if report_type == 'feasibility_study' else ''}
- table_data: list of objects, each with keys {', '.join(table_headers)}
Ensure table_data is realistic, aligns with the project context, and matches the report type.
Use a professional, confident tone. Respond in valid JSON only.
"""
            if self.use_new_client and self.client:
                response = self.client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": "You are a professional analyst. Respond in valid JSON only."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.7,
                    max_tokens=2000
                )
                content = json.loads(response.choices[0].message.content)
                # Validate JSON structure
                required_keys = ['executive_summary', 'recommendations', 'table_data']
                if report_type == 'capital_planning':
                    required_keys.extend(['risk_assessment', 'financial_analysis'])
                elif report_type == 'feasibility_study':
                    required_keys.extend(['risk_assessment', 'findings'])
                elif report_type == 'closeout':
                    required_keys.append('project_review')
                if not all(key in content for key in required_keys):
                    raise ValueError("Invalid JSON structure from AI response")
                # Validate table_data keys
                for row in content['table_data']:
                    if not all(h in row for h in table_headers):
                        raise ValueError("Table row missing required headers")
                return content
            else:
                raise Exception("OpenAI client not initialized.")
        except Exception as e:
            logger.error(f"Error generating AI content: {e}")
            return self._get_fallback_sections(report_type)

    def _get_fallback_sections(self, report_type: str) -> dict:
        template = REPORT_TEMPLATES[report_type]
        table_headers = template['table_headers']
        table_data = [dict(zip(table_headers, row)) for row in template['table_data']]
        fallback = {
            "executive_summary": template['intro'],
            "recommendations": [
                "Prioritize high-impact initiatives.",
                "Monitor risks and adjust plans."
            ],
            "table_data": table_data
        }
        if report_type == 'capital_planning':
            fallback.update({
                "risk_assessment": AI_SECTION_FALLBACKS[report_type],
                "financial_analysis": "Financial projections align with industry standards."
            })
        elif report_type == 'feasibility_study':
            fallback.update({
                "risk_assessment": AI_SECTION_FALLBACKS[report_type],
                "findings": "Project is viable with moderate risk."
            })
        elif report_type == 'closeout':
            fallback.update({
                "project_review": AI_SECTION_FALLBACKS[report_type]
            })
        return fallback

def insert_simple_field(paragraph, field_code):
    """Insert a field into the given paragraph."""
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), field_code)
    paragraph._p.append(fld)
    return fld

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to the paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    run = paragraph.add_run()
    run.text = text
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    run.font.underline = True
    return run

def create_ai_generated_report(company_name: str, industry: str, budget_range: str, 
                              focus_areas: List[str], include_ai_analysis: bool = True) -> str:
    """
    Create an AI-generated capital planning report
    """
    doc = Document()

    # Initialize AI generator
    ai_generator = AIReportGenerator()

    # Generate content
    if include_ai_analysis:
        content = ai_generator.generate_report_content(company_name, industry, budget_range, focus_areas)
    else:
        content = ai_generator._get_fallback_content(company_name, industry, budget_range, focus_areas)

    # --- Page Setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # --- Header Setup ---
    header = section.header
    header.is_linked_to_previous = False

    avail_width = section.page_width - section.left_margin - section.right_margin
    hdr_tbl = header.add_table(rows=1, cols=2, width=avail_width)
    hdr_tbl.autofit = False
    hdr_tbl.columns[0].width = Inches(1.5)
    hdr_tbl.columns[1].width = avail_width - Inches(1.5)

    # Header content
    cell0 = hdr_tbl.rows[0].cells[0]
    p0 = cell0.paragraphs[0]
    p0.add_run(f"{company_name}")

    cell1 = hdr_tbl.rows[0].cells[1]
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run1 = p1.add_run("Page ")
    run1.font.name = 'Arial'
    run1.font.size = Pt(10)
    insert_simple_field(p1, 'PAGE')

    # --- Footer Setup ---
    footer = section.footer
    f_par = footer.paragraphs[0]
    f_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_f = f_par.add_run(f"{company_name} — Confidential | ")
    run_f.font.name = 'Arial'
    run_f.font.size = Pt(9)
    run_f.font.italic = True
    insert_simple_field(f_par, 'DATE \\@ "MMMM d, yyyy"')

    # --- Custom Styles ---
    styles = doc.styles

    # Title Style
    title_style = styles.add_style('EngTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x00, 0x2E, 0x5E)
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_style.paragraph_format.space_after = Pt(12)

    # Subtitle Style
    subtitle_style = styles.add_style('EngSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_style.font.name = 'Arial'
    subtitle_style.font.size = Pt(14)
    subtitle_style.font.bold = True
    subtitle_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    subtitle_style.paragraph_format.space_after = Pt(8)

    # Body Style
    body_style = styles.add_style('EngBody', WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.name = 'Arial'
    body_style.font.size = Pt(11)
    body_style.paragraph_format.line_spacing = 1.15
    body_style.paragraph_format.space_after = Pt(6)

    # Table Header Style
    table_header_style = styles.add_style('EngTableHeader', WD_STYLE_TYPE.PARAGRAPH)
    table_header_style.font.name = 'Arial'
    table_header_style.font.size = Pt(11)
    table_header_style.font.bold = True
    table_header_style.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    table_header_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Document Content ---
    title_p = doc.add_paragraph('Capital Planning Report')
    title_p.style = 'EngTitle'

    # Company Info
    company_info = doc.add_paragraph()
    company_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    company_info.add_run(f"Prepared for: {company_name}")
    company_info.add_run(f" | Industry: {industry}")
    company_info.add_run(f" | Budget Range: {budget_range}")

    # Date
    date_p = doc.add_paragraph()
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run_date = date_p.add_run()
    run_date.font.name = 'Arial'
    run_date.font.size = Pt(11)
    insert_simple_field(date_p, 'DATE \\@ "MMMM d, yyyy"')

    # Executive Summary
    doc.add_paragraph('Executive Summary', style='EngSubtitle')
    p = doc.add_paragraph(content['executive_summary'], style='EngBody')

    # Key Projects
    doc.add_paragraph('Key Capital Projects', style='EngSubtitle')
    for project in content['key_projects']:
        p = doc.add_paragraph(style='EngBody')
        p.add_run(f"{project['name']}: {project['description']}")

    # Capital Expenditure Table
    doc.add_paragraph('Capital Expenditure Overview', style='EngSubtitle')
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr_cells = table.rows[0].cells
    headers = ['Project', 'Budget', 'Timeline', 'ROI (%)']
    for i, hdr in enumerate(headers):
        cell_p = hdr_cells[i].paragraphs[0]
        cell_p.text = hdr
        cell_p.style = 'EngTableHeader'
        hdr_cells[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for project in content['key_projects']:
        row = table.add_row().cells
        row[0].text = project['name']
        row[1].text = project['budget']
        row[2].text = project['timeline']
        row[3].text = project['roi']
        for cell in row:
            cell.paragraphs[0].style = 'EngBody'
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Recommendations
    doc.add_paragraph('Strategic Recommendations', style='EngSubtitle')
    for rec in content['recommendations']:
        p = doc.add_paragraph(rec, style='List Bullet')
        p.style = 'EngBody'

    # Risk Assessment
    doc.add_paragraph('Risk Assessment', style='EngSubtitle')
    doc.add_paragraph(content['risk_assessment'], style='EngBody')

    # Financial Analysis
    doc.add_paragraph('Financial Analysis', style='EngSubtitle')
    doc.add_paragraph(content['financial_analysis'], style='EngBody')

    # AI Analysis Note
    if include_ai_analysis:
        doc.add_paragraph('AI Analysis Note', style='EngSubtitle')
        ai_note = doc.add_paragraph(style='EngBody')
        ai_note.add_run("This report was generated using AI analysis to provide insights and recommendations based on industry best practices and current market trends.")

    # Save document
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f'capital_planning_report_{company_name.replace(" ", "_")}_{timestamp}.docx'
    doc.save(filename)

    return filename

def generate_sample_reports():
    """Generate sample reports for demonstration"""
    sample_data = [
        {
            "company_name": "TechCorp Solutions",
            "industry": "Technology",
            "budget_range": "$5M - $10M",
            "focus_areas": ["AI/ML Infrastructure", "Cloud Migration", "Cybersecurity"]
        },
        {
            "company_name": "GreenEnergy Co",
            "industry": "Renewable Energy",
            "budget_range": "$10M - $25M",
            "focus_areas": ["Solar Infrastructure", "Energy Storage", "Grid Integration"]
        },
        {
            "company_name": "Manufacturing Plus",
            "industry": "Manufacturing",
            "budget_range": "$2M - $5M",
            "focus_areas": ["Automation", "Quality Control", "Supply Chain"]
        }
    ]

    generated_files = []
    for data in sample_data:
        filename = create_ai_generated_report(
            data["company_name"],
            data["industry"],
            data["budget_range"],
            data["focus_areas"],
            include_ai_analysis=True
        )
        generated_files.append({
            "filename": filename,
            "company": data["company_name"],
            "industry": data["industry"]
        })

    return generated_files

REPORT_TEMPLATES = {
    "capital_planning": {
        "title": "VibeOps Capital Planning Report",
        "intro": (
            "This Capital Planning Report outlines the strategic allocation of resources for VibeOps client projects. "
            "It includes detailed analyses of proposed capital expenditures, project timelines, and expected returns on investment."
        ),
        "table_headers": ['Project', 'Budget (USD)', 'Timeline', 'ROI (%)'],
        "table_data": [
            ('Facility Expansion', '$5,000,000', 'Q2 2025 - Q4 2026', '12'),
            ('Equipment Modernization', '$2,500,000', 'Q3 2025 - Q1 2026', '8'),
            ('Technology Integration', '$1,800,000', 'Q4 2025 - Q2 2026', '15'),
            ('Sustainability Initiatives', '$1,200,000', 'Q1 2025 - Q3 2025', '10'),
        ],
        "static_sections": [
            ("Key Capital Projects", [
                'Facility Expansion: New manufacturing plant to increase production capacity by 20%.',
                'Equipment Modernization: Upgrade to energy-efficient machinery to reduce operational costs.',
                'Technology Integration: Implementation of IoT systems for real-time asset monitoring.',
                'Sustainability Initiatives: Installation of solar panels to achieve 30% energy self-sufficiency.'
            ]),
        ],
        "ai_section_title": "AI Analysis & Recommendations",
        "ai_prompt": (
            "You are a VibeOps capital planning expert. Write a detailed, professional analysis and recommendations section for a capital planning report. "
            "Focus on the provided company and project context. Use a confident, expert, and helpful tone."
        ),
    },
    "feasibility_study": {
        "title": "VibeOps Feasibility Study Report",
        "intro": (
            "This Feasibility Study Report evaluates the technical and financial viability of the proposed project for VibeOps clients. "
            "It includes a summary of objectives, constraints, and a professional risk assessment."
        ),
        "table_headers": ['Scenario', 'Estimated Cost', 'Timeline', 'Risk Level'],
        "table_data": [
            ('Base Case', '$2,000,000', '12 months', 'Medium'),
            ('Optimistic', '$1,800,000', '10 months', 'Low'),
            ('Pessimistic', '$2,400,000', '15 months', 'High'),
        ],
        "static_sections": [
            ("Project Objectives", [
                'Deliver scalable infrastructure for future growth.',
                'Minimize operational risk and maximize ROI.',
                'Ensure compliance with all regulatory requirements.'
            ]),
        ],
        "ai_section_title": "AI Findings & Risk Assessment",
        "ai_prompt": (
            "You are a VibeOps feasibility study expert. Write a detailed findings and risk assessment section for a feasibility study report. "
            "Focus on the provided company and project context. Use a confident, expert, and helpful tone."
        ),
    },
    "closeout": {
        "title": "VibeOps Project Closeout Report",
        "intro": (
            "This Project Closeout Report summarizes the completion of the VibeOps client project, including deliverables, lessons learned, and final recommendations."
        ),
        "table_headers": ['Milestone', 'Completion Date', 'Status', 'Notes'],
        "table_data": [
            ('Site Preparation', '2025-03-01', 'Complete', 'No issues'),
            ('Construction', '2025-09-15', 'Complete', 'Minor delays'),
            ('Commissioning', '2025-10-10', 'Complete', 'Successful'),
            ('Handover', '2025-10-20', 'Complete', 'Client satisfied'),
        ],
        "static_sections": [
            ("Lessons Learned", [
                'Early stakeholder engagement reduced change orders.',
                'Regular site meetings improved communication.',
                'Contingency planning mitigated supply chain risks.'
            ]),
        ],
        "ai_section_title": "AI Project Review & Recommendations",
        "ai_prompt": (
            "You are a VibeOps project closeout expert. Write a detailed project review and recommendations section for a closeout report. "
            "Focus on the provided company and project context. Use a confident, expert, and helpful tone."
        ),
    },
}

REPORT_TYPE_CHOICES = [
    ("capital_planning", "Capital Planning Report"),
    ("feasibility_study", "Feasibility Study Report"),
    ("closeout", "Project Closeout Report"),
]

def insert_logos(doc):
    for logo in (os.path.join('static', 'Logo-blk.png'), os.path.join('static', 'Logo-wht.png')):
        try:
            pic = doc.add_picture(logo, width=Inches(1.5))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except Exception:
            continue

# Update ai_generate_section to use AI for the main body sections
ai_generator = AIReportGenerator()
def ai_generate_section(prompt: str, company_name: str, project_context: str, report_type: str = None) -> dict:
    if report_type:
        return ai_generator.generate_report_sections(report_type, company_name, project_context)
    return {}

# Backup AI section text for each template
AI_SECTION_FALLBACKS = {
    "capital_planning": (
        "This section would normally contain a detailed AI-generated analysis and recommendations. "
        "Due to a technical issue, here is a professional fallback: "
        "\n\nVibeOps recommends prioritizing projects with the highest ROI and strategic alignment. "
        "A phased approach to capital deployment will help manage risk and optimize cash flow. "
        "Regularly review project milestones and adjust resource allocation as needed. "
        "Engage stakeholders early and often to ensure alignment and minimize disruptions."
    ),
    "feasibility_study": (
        "This section would normally contain a detailed AI-generated findings and risk assessment. "
        "Due to a technical issue, here is a professional fallback: "
        "\n\nThe project is technically feasible with moderate risk. "
        "Cost estimates are within industry norms, and the timeline is achievable with proper planning. "
        "Key risks include supply chain delays and regulatory changes. "
        "Mitigation strategies should include contingency planning and regular risk reviews."
    ),
    "closeout": (
        "This section would normally contain a detailed AI-generated project review and recommendations. "
        "Due to a technical issue, here is a professional fallback: "
        "\n\nThe project was completed on schedule and within budget. "
        "Lessons learned include the value of early stakeholder engagement and robust risk management. "
        "For future projects, VibeOps recommends enhanced documentation and post-project reviews to capture best practices."
    ),
}

def create_vibeops_report(report_type: str, company_name: str, project_context: str = "") -> str:
    template = REPORT_TEMPLATES[report_type]
    doc = Document()
    # --- Page Setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    # --- Header ---
    header = section.header
    header.is_linked_to_previous = False
    avail_width = section.page_width - section.left_margin - section.right_margin
    hdr_tbl = header.add_table(rows=1, cols=2, width=avail_width)
    hdr_tbl.autofit = False
    hdr_tbl.columns[0].width = Inches(1.5)
    hdr_tbl.columns[1].width = avail_width - Inches(1.5)
    cell0 = hdr_tbl.rows[0].cells[0]
    p0 = cell0.paragraphs[0]
    try:
        run0 = p0.add_run()
        run0.add_picture(os.path.join('static', 'Logo-blk.png'), width=Inches(1.5))
    except Exception:
        p0.add_run("VibeOps Logo")
    cell1 = hdr_tbl.rows[0].cells[1]
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run1 = p1.add_run("Page ")
    run1.font.name = 'Arial'
    run1.font.size = Pt(10)
    insert_simple_field(p1, 'PAGE')
    # --- Footer ---
    footer = section.footer
    f_par = footer.paragraphs[0]
    f_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_f = f_par.add_run("VibeOps — Confidential | ")
    run_f.font.name = 'Arial'
    run_f.font.size = Pt(9)
    run_f.font.italic = True
    insert_simple_field(f_par, 'DATE \\@ "MMMM d, yyyy"')
    try:
        run_logo = f_par.add_run()
        run_logo.add_picture(os.path.join('static', 'Logo-wht.png'), width=Inches(0.75))
    except Exception:
        pass
    # Add a subtle footer line
    f_line = footer.add_paragraph()
    f_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    f_line_run = f_line.add_run("―" * 40)
    f_line_run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
    f_line_run.font.size = Pt(8)
    # --- Custom Styles ---
    styles = doc.styles
    title_style = styles.add_style('EngTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x00, 0x2E, 0x5E)
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_style.paragraph_format.space_after = Pt(18)
    subtitle_style = styles.add_style('EngSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_style.font.name = 'Arial'
    subtitle_style.font.size = Pt(16)
    subtitle_style.font.bold = True
    subtitle_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    subtitle_style.paragraph_format.space_after = Pt(10)
    body_style = styles.add_style('EngBody', WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.name = 'Arial'
    body_style.font.size = Pt(11)
    body_style.paragraph_format.line_spacing = 1.25
    body_style.paragraph_format.space_after = Pt(10)
    table_header_style = styles.add_style('EngTableHeader', WD_STYLE_TYPE.PARAGRAPH)
    table_header_style.font.name = 'Arial'
    table_header_style.font.size = Pt(12)
    table_header_style.font.bold = True
    table_header_style.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    table_header_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # --- Cover Page ---
    title_p = doc.add_paragraph(template['title'])
    title_p.style = 'EngTitle'
    # Horizontal line under title
    hr = doc.add_paragraph()
    hr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hr_run = hr.add_run("―" * 30)
    hr_run.font.color.rgb = RGBColor(0x00, 0x2E, 0x5E)
    hr_run.font.size = Pt(14)
    # Company Info (centered, spaced)
    company_info = doc.add_paragraph()
    company_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    company_info.add_run(f"Prepared for: {company_name}\n").bold = True
    company_info.add_run(f"Date: ")
    insert_simple_field(company_info, 'DATE \\@ "MMMM d, yyyy"')
    company_info.add_run(f"\nProject Context: {project_context}")
    company_info.paragraph_format.space_after = Pt(18)
    # --- Executive Summary ---
    ai_sections = ai_generate_section(template['ai_prompt'], company_name, project_context, report_type=report_type)
    doc.add_paragraph('Executive Summary', style='EngSubtitle')
    doc.add_paragraph(ai_sections.get('executive_summary', template['intro']), style='EngBody')
    # --- Static Sections ---
    for section_title, bullets in template['static_sections']:
        doc.add_paragraph(section_title, style='EngSubtitle')
        for item in bullets:
            p = doc.add_paragraph(item, style='List Bullet')
            p.style = 'EngBody'
    # --- Table ---
    doc.add_paragraph('Overview Table', style='EngSubtitle')
    table = doc.add_table(rows=1, cols=len(template['table_headers']))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for i, hdr in enumerate(template['table_headers']):
        cell_p = hdr_cells[i].paragraphs[0]
        cell_p.text = hdr
        cell_p.style = 'EngTableHeader'
        for run in cell_p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for idx, row_data in enumerate(ai_sections['table_data']):
        row = table.add_row().cells
        for i, hdr in enumerate(template['table_headers']):
            row[i].text = str(row_data[hdr])
            row[i].paragraphs[0].style = 'EngBody'
            row[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if idx % 2 == 1:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F2F6FA')
                row[i]._tc.get_or_add_tcPr().append(shading_elm)
    # --- AI Section(s) ---
    if report_type == 'capital_planning':
        doc.add_paragraph(template['ai_section_title'], style='EngSubtitle')
        doc.add_paragraph(ai_sections.get('financial_analysis', '') or ai_sections.get('risk_assessment', ''), style='EngBody')
        doc.add_paragraph('Strategic Recommendations', style='EngSubtitle')
        for rec in ai_sections.get('recommendations', []):
            p = doc.add_paragraph(rec, style='List Bullet')
            p.style = 'EngBody'
        doc.add_paragraph('Risk Assessment', style='EngSubtitle')
        doc.add_paragraph(ai_sections.get('risk_assessment', ''), style='EngBody')
    elif report_type == 'feasibility_study':
        doc.add_paragraph(template['ai_section_title'], style='EngSubtitle')
        doc.add_paragraph(ai_sections.get('findings', ''), style='EngBody')
        doc.add_paragraph('Recommendations', style='EngSubtitle')
        for rec in ai_sections.get('recommendations', []):
            p = doc.add_paragraph(rec, style='List Bullet')
            p.style = 'EngBody'
        doc.add_paragraph('Risk Assessment', style='EngSubtitle')
        doc.add_paragraph(ai_sections.get('risk_assessment', ''), style='EngBody')
    elif report_type == 'closeout':
        doc.add_paragraph(template['ai_section_title'], style='EngSubtitle')
        doc.add_paragraph(ai_sections.get('project_review', ''), style='EngBody')
        doc.add_paragraph('Recommendations', style='EngSubtitle')
        for rec in ai_sections.get('recommendations', []):
            p = doc.add_paragraph(rec, style='List Bullet')
            p.style = 'EngBody'
    # --- Branding ---
    doc.add_paragraph('Company Branding', style='EngSubtitle')
    insert_logos(doc)
    # --- Hyperlink ---
    p = doc.add_paragraph('For more, visit ', style='EngBody')
    add_hyperlink(p, 'VibeOps.ca', 'https://www.vibeops.ca')
    # --- Appendix ---
    doc.add_page_break()
    doc.add_paragraph('Appendix: Supporting Data', style='EngSubtitle')
    doc.add_paragraph('Detailed cost breakdowns, risk assessments, and technical specifications are available upon request.', style='EngBody')
    # --- Save ---
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f'vibeops_{report_type}_report_{timestamp}.docx'
    doc.save(filename)
    return filename 
