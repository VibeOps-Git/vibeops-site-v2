# dot_print.py
"""
Dot Print — server-side document parser for Reportly.

- parse_docx_structure(path): returns ordered sections (H1/H2/H3) with style names
- extract_plaintext(path, max_chars): returns flattened body text for LLM context
- pdf_to_docx(pdf_path, out_docx): robust PDF→DOCX via LibreOffice (soffice)

This module replaces/supplements the ad-hoc extract_template_structure so both the
UI tool and internal ingestion share the exact same parsing logic.
"""

from pathlib import Path
from docx import Document
import subprocess, shlex, time, os, re
from typing import List, Dict, Any

_WS = re.compile(r"\s+")

def _normalize_ws(s: str) -> str:
    return _WS.sub(" ", (s or "")).strip()

def parse_docx_structure(docx_path: str) -> List[Dict[str, Any]]:
    """
    Extract headings (H1/H2/H3) in order, preserving style names.
    Returns: [{title, level, style_name}, ...]
    """
    doc = Document(str(docx_path))
    sections = []
    for p in doc.paragraphs:
        style_name = (p.style.name if p.style else "") or ""
        text = _normalize_ws(p.text)
        if not text:
            continue
        lower = style_name.lower()
        level = 0
        if "heading 1" in lower:
            level = 1
        elif "heading 2" in lower:
            level = 2
        elif "heading 3" in lower:
            level = 3

        if level:
            sections.append({
                "title": text,
                "level": level,
                "style_name": style_name
            })

    if not sections:
        # Sensible defaults if a template has no explicit headings
        for title in ["Executive Summary", "Scope", "Findings", "Risks", "Recommendations"]:
            sections.append({"title": title, "level": 1, "style_name": "Heading 1"})

    return sections

def extract_plaintext(docx_path: str, max_chars: int = 15000) -> str:
    """Flatten body text to plain text (headers/footers not included)."""
    doc = Document(str(docx_path))
    out, size = [], 0
    for p in doc.paragraphs:
        t = _normalize_ws(p.text)
        if t:
            out.append(t)
            size += len(t) + 1
            if size >= max_chars:
                break
    return "\n".join(out)[:max_chars]

def pdf_to_docx(pdf_path: str, out_docx: str) -> str:
    """
    Convert PDF→DOCX using LibreOffice. Robust to LO’s naming quirks.
    - Writes into the destination directory
    - Picks the most likely produced .docx by stem or newest
    Returns the final docx path.
    """
    pdf_path = Path(pdf_path).resolve()
    out_docx = Path(out_docx).resolve()
    out_dir = out_docx.parent
    out_dir.mkdir(parents=True, exist_ok=True)

    before = {p.name for p in out_dir.glob("*.docx")}
    start = time.time()

    cmd = (
        f'soffice --headless --convert-to docx --outdir {shlex.quote(str(out_dir))} '
        f'{shlex.quote(str(pdf_path))}'
    )
    proc = subprocess.run(cmd, shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    if proc.returncode != 0:
        msg = proc.stderr.decode("utf-8", "ignore") or proc.stdout.decode("utf-8", "ignore")
        raise RuntimeError(f"LibreOffice conversion failed: {msg}")

    # Find produced file(s)
    candidates = []
    for p in out_dir.glob("*.docx"):
        if p.name in before:
            continue
        try:
            mtime = p.stat().st_mtime
        except Exception:
            mtime = 0
        candidates.append((p, mtime))

    if not candidates:
        # Fallback: try by stem match in case timestamps didn’t change
        stem_guess = out_dir / (pdf_path.stem + ".docx")
        if stem_guess.exists():
            produced = stem_guess
        else:
            raise FileNotFoundError("PDF→DOCX produced no .docx output in target dir.")
    else:
        # Prefer exact stem match; else newest
        exact = [p for (p, _) in candidates if p.stem == pdf_path.stem]
        produced = exact[0] if exact else sorted(candidates, key=lambda x: x[1], reverse=True)[0][0]

    if produced.resolve() != out_docx:
        # copy/rename to the requested output filename
        if out_docx.exists():
            out_docx.unlink()
        produced.replace(out_docx)

    return str(out_docx)
