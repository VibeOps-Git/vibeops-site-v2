#!/usr/bin/env python3
"""
ai_report_generator.py

AI-powered capital planning report generator using OpenAI API
"""

import os
import json
import datetime
from typing import Dict, List, Optional
from openai import OpenAI
import logging
import traceback
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
from PIL import Image

# Set OpenAI API key globally

# Set up logger
logger = logging.getLogger("vibeops.ai_report_generator")
logger.setLevel(logging.DEBUG)
if not logger.hasHandlers():
    handler = logging.StreamHandler()
    formatter = logging.Formatter('[%(asctime)s] %(levelname)s %(name)s: %(message)s')
    handler.setFormatter(formatter)
    logger.addHandler(handler)

# from openai import OpenAI
# client = OpenAI(api_key=os.getenv('OpenAI_API_KEY'))

# Restore and update AIReportGenerator
class AIReportGenerator:
    def __init__(self):
        logger.debug("Initializing AIReportGenerator...")
        self.api_key = os.getenv('OpenAI_API_KEY')
        if self.api_key:
            try:
                self.client = OpenAI(api_key=self.api_key)
                self.use_new_client = True
                logger.debug("Using OpenAI client.")
            except Exception as e:
                logger.warning(f"Could not initialize OpenAI client: {e}")
                self.client = None
                self.use_new_client = False
        else:
            logger.warning("No OpenAI API key found. AI features will be disabled.")
            self.client = None
            self.use_new_client = False

    def generate_report_sections(self, report_type: str, company_name: str, project_context: str = "") -> dict:
        logger.debug(f"Generating AI report sections for {report_type}, {company_name}, {project_context}")
        try:
            template = REPORT_TEMPLATES[report_type]
            table_headers = template['table_headers']

            # Determine which extra JSON fields to request
            extra_fields = []
            if report_type == 'defense_compliance':
                extra_fields = ['compliance_summary: string', 'risk_assessment: string']
            elif report_type == 'capital_planning':
                extra_fields = ['risk_assessment: string', 'financial_analysis: string']
            elif report_type == 'feasibility_study':
                extra_fields = ['risk_assessment: string', 'findings: string']
            elif report_type == 'construction_closeout':
                extra_fields = ['project_review: string']

            # Build the prompt
            prompt = f"""
You are a VibeOps {report_type.replace('_', ' ')} expert.
Generate a detailed report for {company_name} with project context: {project_context}.
Return a JSON object with:
- executive_summary: string
- recommendations: list of strings
"""
            for field in extra_fields:
                prompt += f"- {field}\n"
            prompt += f"- table_data: list of objects, each with keys {', '.join(table_headers)}\n\n"
            prompt += "Ensure table_data is realistic, aligns with the project context, and matches the report type.\n"
            prompt += "Use a professional, confident tone. Respond in valid JSON only."

            # Call the OpenAI client
            response = self.client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a professional analyst. Respond in valid JSON only."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=2000
            )

            content = json.loads(response.choices[0].message.content)

            # Validate that all required keys were returned
            required_keys = ['executive_summary', 'recommendations', 'table_data']
            if report_type == 'defense_compliance':
                required_keys += ['compliance_summary', 'risk_assessment']
            elif report_type == 'capital_planning':
                required_keys += ['risk_assessment', 'financial_analysis']
            elif report_type == 'feasibility_study':
                required_keys += ['risk_assessment', 'findings']
            elif report_type == 'construction_closeout':
                required_keys.append('project_review')

            missing = [k for k in required_keys if k not in content]
            if missing:
                raise ValueError(f"Missing keys in AI response: {missing}")

            # Validate each row of table_data
            for row in content['table_data']:
                if not all(h in row for h in table_headers):
                    raise ValueError("Table row missing required headers")

            return content

        except Exception as e:
            logger.error(f"Error generating AI content: {e}")
            return self._get_fallback_sections(report_type)


    def _get_fallback_sections(self, report_type: str) -> dict:
        template = REPORT_TEMPLATES[report_type]
        table_headers = template['table_headers']
        table_data = [dict(zip(table_headers, row)) for row in template['table_data']]
        fallback = {
            "executive_summary": template['intro'],
            "recommendations": [
                "Prioritize high-impact initiatives.",
                "Monitor risks and adjust plans."
            ],
            "table_data": table_data
        }
        if report_type == 'capital_planning':
            fallback.update({
                "risk_assessment": AI_SECTION_FALLBACKS[report_type],
                "financial_analysis": "Financial projections align with industry standards."
            })
        elif report_type == 'feasibility_study':
            fallback.update({
                "risk_assessment": AI_SECTION_FALLBACKS[report_type],
                "findings": "Project is viable with moderate risk."
            })
        elif report_type == 'closeout':
            fallback.update({
                "project_review": AI_SECTION_FALLBACKS[report_type]
            })
        return fallback

def validate_input(text: str, default: str) -> str:
    if not text or not re.match(r'^[a-zA-Z0-9\s\.,\-]+$', text):
        logger.warning(f"Invalid input '{text}', using default: {default}")
        return default
    return text

def format_price(value):
    # Try to format as price if it looks like a number or $-prefixed string
    try:
        if isinstance(value, (int, float)):
            return f"${value:,.0f}"
        s = str(value).replace('$', '').replace(',', '').replace(' ', '')
        if s.replace('.', '', 1).isdigit():
            return f"${float(s):,.0f}"
        # Try to extract numbers from things like '500000' or '500000.00'
        match = re.match(r'\$?([0-9]+(?:\.[0-9]+)?)', str(value))
        if match:
            return f"${float(match.group(1)):,}"
    except Exception:
        pass
    return str(value)

def add_scaled_logo(paragraph, logo_path, max_width=Inches(1.5), max_pixel_height=40):
    try:
        with Image.open(logo_path) as img:
            width, height = img.size
            # Always restrict by pixel height
            scale = min(1.0, max_pixel_height / height)
            width = int(width * scale)
            height = int(height * scale)
            # Convert to inches (assuming 96 DPI)
            width_inches = width / 96
            height_inches = height / 96
            run = paragraph.add_run()
            run.add_picture(logo_path, width=Inches(width_inches), height=Inches(height_inches))
    except Exception as e:
        logger.error(f"Failed to add logo {logo_path}: {e}")
        paragraph.add_run("[Logo]")

def insert_simple_field(paragraph, field_code):
    try:
        fld = OxmlElement('w:fldSimple')
        fld.set(qn('w:instr'), field_code)
        paragraph._p.append(fld)
        return fld
    except Exception as e:
        logger.error(f"Failed to insert field {field_code}: {e}")
        paragraph.add_run(f"[Field Error: {field_code}]")
        return None

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to the paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    run = paragraph.add_run()
    run.text = text
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    run.font.underline = True
    return run

def create_ai_generated_report(report_type, company_name: str, industry: str, budget_range: str, 
                              focus_areas: List[str], include_ai_analysis: bool = True) -> str:
    """
    Create an AI-generated capital planning report
    """
    doc = Document()

    # Initialize AI generator
    ai_generator = AIReportGenerator()

    # Generate content
    if include_ai_analysis:
        content = ai_generator.generate_report_sections(company_name, industry, budget_range, focus_areas)
    else:
        content = ai_generator._get_fallback_content(company_name, industry, budget_range, focus_areas)

    # --- Page Setup ---
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # --- Header Setup ---
    header = section.header
    header.is_linked_to_previous = False

    avail_width = section.page_width - section.left_margin - section.right_margin
    hdr_tbl = header.add_table(rows=1, cols=2, width=avail_width)
    hdr_tbl.autofit = False
    hdr_tbl.columns[0].width = Inches(1.5)
    hdr_tbl.columns[1].width = avail_width - Inches(1.5)

    # Header content
    cell0 = hdr_tbl.rows[0].cells[0]
    p0 = cell0.paragraphs[0]
    p0.add_run(f"{company_name}")

    cell1 = hdr_tbl.rows[0].cells[1]
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run1 = p1.add_run("Page ")
    run1.font.name = 'Arial'
    run1.font.size = Pt(10)
    insert_simple_field(p1, 'PAGE')

    # --- Footer Setup ---
    footer = section.footer
    f_par = footer.paragraphs[0]
    f_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_f = f_par.add_run(f"{company_name} — Confidential | ")
    run_f.font.name = 'Arial'
    run_f.font.size = Pt(9)
    run_f.font.italic = True
    insert_simple_field(f_par, 'DATE \\@ "MMMM d, yyyy"')

    # --- Custom Styles ---
    styles = doc.styles

    # Title Style
    title_style = styles.add_style('EngTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x00, 0x2E, 0x5E)
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_style.paragraph_format.space_after = Pt(12)

    # Subtitle Style
    subtitle_style = styles.add_style('EngSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_style.font.name = 'Arial'
    subtitle_style.font.size = Pt(14)
    subtitle_style.font.bold = True
    subtitle_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    subtitle_style.paragraph_format.space_after = Pt(8)

    # Body Style
    body_style = styles.add_style('EngBody', WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.name = 'Arial'
    body_style.font.size = Pt(11)
    body_style.paragraph_format.line_spacing = 1.15
    body_style.paragraph_format.space_after = Pt(6)

    # Table Header Style
    table_header_style = styles.add_style('EngTableHeader', WD_STYLE_TYPE.PARAGRAPH)
    table_header_style.font.name = 'Arial'
    table_header_style.font.size = Pt(11)
    table_header_style.font.bold = True
    table_header_style.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    table_header_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Document Content ---
    title_p = doc.add_paragraph('Capital Planning Report')
    title_p.style = 'EngTitle'

    # Company Info
    company_info = doc.add_paragraph()
    company_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    company_info.add_run(f"Prepared for: {company_name}")
    company_info.add_run(f" | Industry: {industry}")
    company_info.add_run(f" | Budget Range: {budget_range}")

    # Date
    date_p = doc.add_paragraph()
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run_date = date_p.add_run()
    run_date.font.name = 'Arial'
    run_date.font.size = Pt(11)
    insert_simple_field(date_p, 'DATE \\@ "MMMM d, yyyy"')

    # Executive Summary
    doc.add_paragraph('Executive Summary', style='EngSubtitle')
    p = doc.add_paragraph(content['executive_summary'], style='EngBody')

    # Key Projects
    doc.add_paragraph('Key Capital Projects', style='EngSubtitle')
    for project in content['key_projects']:
        p = doc.add_paragraph(style='EngBody')
        p.add_run(f"{project['name']}: {project['description']}")

    # Capital Expenditure Table
    doc.add_paragraph('Capital Expenditure Overview', style='EngSubtitle')
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr_cells = table.rows[0].cells
    headers = ['Project', 'Budget', 'Timeline', 'ROI (%)']
    for i, hdr in enumerate(headers):
        cell_p = hdr_cells[i].paragraphs[0]
        cell_p.text = hdr
        cell_p.style = 'EngTableHeader'
        hdr_cells[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for project in content['key_projects']:
        row = table.add_row().cells
        row[0].text = project['name']
        row[1].text = project['budget']
        row[2].text = project['timeline']
        row[3].text = project['roi']
        for cell in row:
            cell.paragraphs[0].style = 'EngBody'
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Recommendations
    doc.add_paragraph('Strategic Recommendations', style='EngSubtitle')
    for rec in content['recommendations']:
        p = doc.add_paragraph(rec, style='List Bullet')
        p.style = 'EngBody'

    # Risk Assessment
    doc.add_paragraph('Risk Assessment', style='EngSubtitle')
    doc.add_paragraph(content['risk_assessment'], style='EngBody')

    # Financial Analysis
    doc.add_paragraph('Financial Analysis', style='EngSubtitle')
    doc.add_paragraph(content['financial_analysis'], style='EngBody')

    # AI Analysis Note
    if include_ai_analysis:
        doc.add_paragraph('AI Analysis Note', style='EngSubtitle')
        ai_note = doc.add_paragraph(style='EngBody')
        ai_note.add_run("This report was generated using AI analysis to provide insights and recommendations based on industry best practices and current market trends.")

    # Save document
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f'capital_planning_report_{company_name.replace(" ", "_")}_{timestamp}.docx'
    doc.save(filename)

    return filename

def generate_sample_reports():
    """Generate sample reports for demonstration"""
    sample_data = [
        {
            "company_name": "TechCorp Solutions",
            "industry": "Technology",
            "budget_range": "$5M - $10M",
            "focus_areas": ["AI/ML Infrastructure", "Cloud Migration", "Cybersecurity"]
        },
        {
            "company_name": "GreenEnergy Co",
            "industry": "Renewable Energy",
            "budget_range": "$10M - $25M",
            "focus_areas": ["Solar Infrastructure", "Energy Storage", "Grid Integration"]
        },
        {
            "company_name": "Manufacturing Plus",
            "industry": "Manufacturing",
            "budget_range": "$2M - $5M",
            "focus_areas": ["Automation", "Quality Control", "Supply Chain"]
        }
    ]

    generated_files = []
    for data in sample_data:
        filename = create_ai_generated_report(
            data["company_name"],
            data["industry"],
            data["budget_range"],
            data["focus_areas"],
            include_ai_analysis=True
        )
        generated_files.append({
            "filename": filename,
            "company": data["company_name"],
            "industry": data["industry"]
        })

    return generated_files

REPORT_TEMPLATES = {
    "engineering_feasibility": {
        "title": "VibeOps Engineering Feasibility Study Report",
        "intro": (
            "This Engineering Feasibility Study evaluates the technical and economic viability of the proposed engineering project for VibeOps clients. "
            "It provides a comprehensive analysis of project objectives, constraints, and risks to ensure informed decision-making."
        ),
        "table_headers": ["Project Component", "Estimated Cost", "Timeline", "Risk Level"],
        "table_data": [
            ("Structural Design", "$1,500,000", "6 months", "Low"),
            ("Material Procurement", "$800,000", "3 months", "Medium"),
            ("Site Preparation", "$600,000", "2 months", "Low"),
        ],
        "static_sections": [
            ("Project Objectives", [
                "Validate technical feasibility of proposed engineering solutions.",
                "Optimize cost efficiency while maintaining quality standards.",
                "Ensure compliance with engineering regulations and standards."
            ]),
        ],
        "ai_section_title": "AI Findings & Risk Assessment",
        "ai_prompt": (
            "You are a VibeOps engineering feasibility expert. Generate a detailed feasibility study report for {company_name} with project context: {project_context}. "
            "Return a JSON object with: "
            "- executive_summary: string summarizing the feasibility study "
            "- recommendations: list of actionable recommendations "
            "- findings: string detailing technical and economic findings "
            "- risk_assessment: string analyzing potential risks "
            "- table_data: list of objects with keys Project Component, Estimated Cost, Timeline, Risk Level "
            "Ensure table_data is realistic, aligns with the project context, and matches engineering feasibility requirements. "
            "Use a professional, confident tone. Respond in valid JSON only."
        ),
        "fallback_sections": {
            "executive_summary": "This section would normally contain a detailed AI-generated executive summary. Due to a technical issue, here is a professional fallback.",
            "recommendations": [
                "Prioritize projects with the highest ROI and strategic alignment.",
                "A phased approach to capital deployment will help manage risk and optimize cash flow.",
                "Regularly review project milestones and adjust resource allocation as needed.",
                "Engage stakeholders early and often to ensure alignment and minimize disruptions."
            ],
            "findings": "This section would normally contain a detailed AI-generated findings and risk assessment. Due to a technical issue, here is a professional fallback.",
            "risk_assessment": "This section would normally contain a detailed AI-generated risk assessment. Due to a technical issue, here is a professional fallback.",
            "financial_analysis": "This section would normally contain a detailed AI-generated financial analysis. Due to a technical issue, here is a professional fallback.",
            "project_review": "This section would normally contain a detailed AI-generated project review and recommendations. Due to a technical issue, here is a professional fallback.",
        }
    },
    "defense_compliance": {
        "title": "VibeOps Defense Compliance Report",
        "intro": (
            "This Defense Compliance Report ensures that the project adheres to strict regulatory and security standards required for defense contracts. "
            "It outlines compliance measures, audit readiness, and risk mitigation strategies."
        ),
        "table_headers": ["Compliance Area", "Status", "Audit Date", "Notes"],
        "table_data": [
            ("Security Protocols", "Compliant", "2025-06-15", "Meets DoD standards"),
            ("Data Protection", "In Progress", "2025-07-01", "Requires encryption update"),
            ("Personnel Clearance", "Compliant", "2025-06-01", "All staff cleared"),
        ],
        "static_sections": [
            ("Compliance Objectives", [
                "Ensure adherence to DoD and federal regulations.",
                "Maintain audit-ready documentation for all compliance areas.",
                "Implement robust cybersecurity measures."
            ]),
        ],
        "ai_section_title": "AI Compliance Analysis & Recommendations",
        "ai_prompt": (
            "You are a VibeOps defense compliance expert. Generate a detailed compliance report for {company_name} with project context: {project_context}. "
            "Return a JSON object with: "
            "- executive_summary: string summarizing compliance status "
            "- recommendations: list of actionable recommendations "
            "- compliance_summary: string detailing compliance status and audit readiness "
            "- risk_assessment: string analyzing compliance risks "
            "- table_data: list of objects with keys Compliance Area, Status, Audit Date, Notes "
            "Ensure table_data is realistic, aligns with defense contract requirements, and matches the project context. "
            "Use a professional, confident tone. Respond in valid JSON only."
        ),
        "fallback_sections": {
            "executive_summary": "This section would normally contain a detailed AI-generated executive summary. Due to a technical issue, here is a professional fallback.",
            "recommendations": [
                "Ensure adherence to DoD and federal regulations.",
                "Maintain audit-ready documentation for all compliance areas.",
                "Implement robust cybersecurity measures."
            ],
            "compliance_summary": "This section would normally contain a detailed AI-generated compliance summary. Due to a technical issue, here is a professional fallback.",
            "risk_assessment": "This section would normally contain a detailed AI-generated risk assessment. Due to a technical issue, here is a professional fallback.",
            "project_review": "This section would normally contain a detailed AI-generated project review and recommendations. Due to a technical issue, here is a professional fallback.",
        }
    },
    "utilities_estimate": {
        "title": "VibeOps Utilities Infrastructure Estimate Report",
        "intro": (
            "This Utilities Infrastructure Estimate Report provides detailed cost and material estimates for pipelines, grids, or network projects. "
            "It includes a bill of materials (BOM), timelines, and strategic recommendations."
        ),
        "table_headers": ["Component", "Quantity", "Unit Cost", "Total Cost"],
        "table_data": [
            ("Pipeline Segment", "10 km", "$50,000/km", "$500,000"),
            ("Substation Upgrade", "2 units", "$200,000/unit", "$400,000"),
            ("Control Systems", "5 units", "$30,000/unit", "$150,000"),
        ],
        "static_sections": [
            ("Project Scope", [
                "Develop accurate cost estimates for infrastructure components.",
                "Ensure scalability for future utility expansions.",
                "Minimize environmental impact during implementation."
            ]),
        ],
        "ai_section_title": "AI Cost Analysis & Recommendations",
        "ai_prompt": (
            "You are a VibeOps utilities infrastructure expert. Generate a detailed estimate report for {company_name} with project context: {project_context}. "
            "Return a JSON object with: "
            "- executive_summary: string summarizing the cost estimate "
            "- recommendations: list of actionable recommendations "
            "- cost_breakdown: string detailing cost analysis and BOM "
            "- table_data: list of objects with keys Component, Quantity, Unit Cost, Total Cost "
            "Ensure table_data is realistic, aligns with utilities infrastructure requirements, and matches the project context. "
            "Use a professional, confident tone. Respond in valid JSON only."
        ),
        "fallback_sections": {
            "executive_summary": "This section would normally contain a detailed AI-generated executive summary. Due to a technical issue, here is a professional fallback.",
            "recommendations": [
                "Develop accurate cost estimates for infrastructure components.",
                "Ensure scalability for future utility expansions.",
                "Minimize environmental impact during implementation."
            ],
            "cost_breakdown": "This section would normally contain a detailed AI-generated cost breakdown. Due to a technical issue, here is a professional fallback.",
            "project_review": "This section would normally contain a detailed AI-generated project review and recommendations. Due to a technical issue, here is a professional fallback.",
        }
    },
    "construction_closeout": {
        "title": "VibeOps Construction Closeout & Handover Report",
        "intro": (
            "This Construction Closeout & Handover Report summarizes the completion of the construction project, including deliverables, lessons learned, and final recommendations."
        ),
        "table_headers": ["Milestone", "Completion Date", "Status", "Notes"],
        "table_data": [
            ("Foundation Work", "2025-04-01", "Complete", "No issues"),
            ("Structural Build", "2025-08-15", "Complete", "Minor delays"),
            ("Final Inspections", "2025-09-30", "Complete", "Passed"),
        ],
        "static_sections": [
            ("Lessons Learned", [
                "Early coordination with subcontractors reduced delays.",
                "Regular quality checks improved outcomes.",
                "Documentation streamlined handover process."
            ]),
        ],
        "ai_section_title": "AI Project Review & Recommendations",
        "ai_prompt": (
            "You are a VibeOps construction closeout expert. Generate a detailed closeout report for {company_name} with project context: {project_context}. "
            "Return a JSON object with: "
            "- executive_summary: string summarizing project completion "
            "- recommendations: list of actionable recommendations "
            "- project_review: string detailing project outcomes and lessons learned "
            "- table_data: list of objects with keys Milestone, Completion Date, Status, Notes "
            "Ensure table_data is realistic, aligns with construction project requirements, and matches the project context. "
            "Use a professional, confident tone. Respond in valid JSON only."
        ),
        "fallback_sections": {
            "executive_summary": "This section would normally contain a detailed AI-generated executive summary. Due to a technical issue, here is a professional fallback.",
            "recommendations": [
                "The project was completed on schedule and within budget.",
                "Lessons learned include the value of early stakeholder engagement and robust risk management.",
                "For future projects, VibeOps recommends enhanced documentation and post-project reviews to capture best practices."
            ],
            "project_review": "This section would normally contain a detailed AI-generated project review. Due to a technical issue, here is a professional fallback.",
        }
    },
}

REPORT_TYPE_CHOICES = [
    ("engineering_feasibility", "Engineering Feasibility Study"),
    ("defense_compliance", "Defense Compliance Report"),
    ("utilities_estimate", "Utilities Infrastructure Estimate"),
    ("construction_closeout", "Construction Closeout & Handover"),
]

def insert_logos(doc):
    for logo in (os.path.join('static', 'Logo-blk.png'), os.path.join('static', 'Logo-wht.png')):
        try:
            pic = doc.add_picture(logo, width=Inches(1.5))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except Exception:
            continue

# Update ai_generate_section to use AI for the main body sections
ai_generator = AIReportGenerator()
def ai_generate_section(prompt: str, company_name: str, project_context: str, report_type: str = None) -> dict:
    if report_type:
        return ai_generator.generate_report_sections(report_type, company_name, project_context)
    return {}

# Backup AI section text for each template
AI_SECTION_FALLBACKS = {
    "capital_planning": (
        "This section would normally contain a detailed AI-generated analysis and recommendations. "
        "Due to a technical issue, here is a professional fallback: "
        "\n\nVibeOps recommends prioritizing projects with the highest ROI and strategic alignment. "
        "A phased approach to capital deployment will help manage risk and optimize cash flow. "
        "Regularly review project milestones and adjust resource allocation as needed. "
        "Engage stakeholders early and often to ensure alignment and minimize disruptions."
    ),
    "feasibility_study": (
        "This section would normally contain a detailed AI-generated findings and risk assessment. "
        "Due to a technical issue, here is a professional fallback: "
        "\n\nThe project is technically feasible with moderate risk. "
        "Cost estimates are within industry norms, and the timeline is achievable with proper planning. "
        "Key risks include supply chain delays and regulatory changes. "
        "Mitigation strategies should include contingency planning and regular risk reviews."
    ),
    "closeout": (
        "This section would normally contain a detailed AI-generated project review and recommendations. "
        "Due to a technical issue, here is a professional fallback: "
        "\n\nThe project was completed on schedule and within budget. "
        "Lessons learned include the value of early stakeholder engagement and robust risk management. "
        "For future projects, VibeOps recommends enhanced documentation and post-project reviews to capture best practices."
    ),
}

def create_vibeops_report(report_type: str, company_name: str, project_context: str = "", logo_path: str = None) -> str:
    template = REPORT_TEMPLATES[report_type]
    company_name_input = company_name
    company_name = validate_input(company_name, "Unknown Company")
    project_context = validate_input(project_context, "General Project")
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    # --- Header ---
    header = section.header
    header.is_linked_to_previous = False
    avail_width = section.page_width - section.left_margin - section.right_margin
    hdr_tbl = header.add_table(rows=1, cols=2, width=avail_width)
    hdr_tbl.autofit = False
    hdr_tbl.columns[0].width = Inches(1.5)
    hdr_tbl.columns[1].width = avail_width - Inches(1.5)
    cell0 = hdr_tbl.rows[0].cells[0]
    # Always use a new paragraph for the logo, and keep text separate
    p0 = cell0.paragraphs[0]
    p0.clear()
    if logo_path:
        add_scaled_logo(p0, logo_path, max_width=Inches(1.5), max_pixel_height=40)
    else:
        add_scaled_logo(p0, os.path.join('static', 'Logo-blk.png'), max_width=Inches(1.5), max_pixel_height=40)
    cell1 = hdr_tbl.rows[0].cells[1]
    p1 = cell1.paragraphs[0]
    p1.clear()
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run1 = p1.add_run(f"{company_name_input.strip() if company_name_input.strip() else 'Your Company'} Report  |  Page ")
    run1.font.name = 'Arial'
    run1.font.size = Pt(10)
    insert_simple_field(p1, 'PAGE')
    # --- Footer ---
    footer = section.footer
    f_par = footer.add_paragraph()
    f_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_company = company_name_input.strip() if company_name_input.strip() else "Your Company"
    run_f = f_par.add_run(f"{footer_company} — Confidential | ")
    run_f.font.name = 'Arial'
    run_f.font.size = Pt(9)
    run_f.font.italic = True
    insert_simple_field(f_par, 'DATE \\@ "MMMM d, yyyy"')
    # Always use a new paragraph for the logo in the footer
    logo_footer_par = footer.add_paragraph()
    logo_footer_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    try:
        if logo_path:
            add_scaled_logo(logo_footer_par, logo_path, max_width=Inches(0.75), max_pixel_height=30)
        else:
            add_scaled_logo(logo_footer_par, os.path.join('static', 'Logo-wht.png'), max_width=Inches(0.75), max_pixel_height=30)
    except Exception:
        pass
    # Add a subtle footer line
    f_line = footer.add_paragraph()
    f_line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    f_line_run = f_line.add_run("―" * 40)
    f_line_run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
    f_line_run.font.size = Pt(8)
    # --- Custom Styles ---
    styles = doc.styles
    title_style = styles.add_style('EngTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x00, 0x2E, 0x5E)
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_style.paragraph_format.space_after = Pt(18)
    subtitle_style = styles.add_style('EngSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_style.font.name = 'Arial'
    subtitle_style.font.size = Pt(16)
    subtitle_style.font.bold = True
    subtitle_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    subtitle_style.paragraph_format.space_after = Pt(10)
    body_style = styles.add_style('EngBody', WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.name = 'Arial'
    body_style.font.size = Pt(11)
    body_style.paragraph_format.line_spacing = 1.25
    body_style.paragraph_format.space_after = Pt(10)
    table_header_style = styles.add_style('EngTableHeader', WD_STYLE_TYPE.PARAGRAPH)
    table_header_style.font.name = 'Arial'
    table_header_style.font.size = Pt(12)
    table_header_style.font.bold = True
    table_header_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black text
    table_header_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # --- Cover Page ---
    title_p = doc.add_paragraph(template['title'])
    title_p.style = 'EngTitle'
    hr = doc.add_paragraph()
    hr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hr_run = hr.add_run("―" * 30)
    hr_run.font.color.rgb = RGBColor(0x00, 0x2E, 0x5E)
    hr_run.font.size = Pt(14)
    company_info = doc.add_paragraph()
    company_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    company_info.add_run(f"Prepared for: {company_name}\n").bold = True
    company_info.add_run(f"Date: ")
    insert_simple_field(company_info, 'DATE \\@ "MMMM d, yyyy"')
    company_info.add_run(f"\nProject Context: {project_context}")
    company_info.paragraph_format.space_after = Pt(18)
    # --- AI Section Handling ---
    ai_sections = ai_generate_section(template['ai_prompt'], company_name, project_context, report_type=report_type)
    required_keys = ['executive_summary', 'recommendations', 'table_data']
    if report_type == 'defense_compliance':
        required_keys += ['compliance_summary', 'risk_assessment']
    elif report_type == 'engineering_feasibility':
        required_keys += ['findings', 'risk_assessment']
    elif report_type == 'utilities_estimate':
        required_keys += ['cost_breakdown', 'project_scope']
    elif report_type == 'construction_closeout':
        required_keys += ['project_review']
    fallback_sections = template.get('fallback_sections', {})
    for key in required_keys:
        if key not in ai_sections:
            logger.error(f"Missing AI section: {key}")
            if key == 'cost_breakdown' and 'table_data' in ai_sections:
                try:
                    total = 0
                    for row in ai_sections['table_data']:
                        val = row.get('Total Cost') or row.get('Estimated Cost') or row.get('Budget (USD)')
                        if val:
                            try:
                                total += float(str(val).replace('$','').replace(',','').replace(' ','').replace('USD',''))
                            except Exception:
                                pass
                    ai_sections[key] = f"Total estimated cost: ${total:,.0f}" if total else fallback_sections.get(key, "")
                except Exception:
                    ai_sections[key] = fallback_sections.get(key, "")
            else:
                ai_sections[key] = fallback_sections.get(key, "")
    # --- Executive Summary ---
    doc.add_paragraph('Executive Summary', style='EngSubtitle')
    doc.add_paragraph(ai_sections.get('executive_summary', template['intro']), style='EngBody')
    # --- Template-specific AI Sections ---
    if report_type == 'defense_compliance':
        doc.add_paragraph('Compliance Summary', style='EngSubtitle')
        doc.add_paragraph(ai_sections.get('compliance_summary', ''), style='EngBody')
        doc.add_paragraph('Risk Assessment', style='EngSubtitle')
        doc.add_paragraph(ai_sections.get('risk_assessment', ''), style='EngBody')
        doc.add_paragraph('Recommendations', style='EngSubtitle')
        for rec in ai_sections.get('recommendations', []):
            p = doc.add_paragraph(rec, style='List Bullet')
            p.style = 'EngBody'
    elif report_type == 'engineering_feasibility':
        doc.add_paragraph('Findings', style='EngSubtitle')
        doc.add_paragraph(ai_sections.get('findings', ''), style='EngBody')
        doc.add_paragraph('Risk Assessment', style='EngSubtitle')
        doc.add_paragraph(ai_sections.get('risk_assessment', ''), style='EngBody')
        doc.add_paragraph('Recommendations', style='EngSubtitle')
        for rec in ai_sections.get('recommendations', []):
            p = doc.add_paragraph(rec, style='List Bullet')
            p.style = 'EngBody'
    elif report_type == 'utilities_estimate':
        doc.add_paragraph('Project Scope', style='EngSubtitle')
        doc.add_paragraph(ai_sections.get('project_scope', ''), style='EngBody')
        doc.add_paragraph('Cost Breakdown', style='EngSubtitle')
        doc.add_paragraph(ai_sections.get('cost_breakdown', ''), style='EngBody')
        doc.add_paragraph('Recommendations', style='EngSubtitle')
        for rec in ai_sections.get('recommendations', []):
            p = doc.add_paragraph(rec, style='List Bullet')
            p.style = 'EngBody'
    elif report_type == 'construction_closeout':
        doc.add_paragraph('Project Review', style='EngSubtitle')
        doc.add_paragraph(ai_sections.get('project_review', ''), style='EngBody')
        doc.add_paragraph('Recommendations', style='EngSubtitle')
        for rec in ai_sections.get('recommendations', []):
            p = doc.add_paragraph(rec, style='List Bullet')
            p.style = 'EngBody'
    # --- Static Sections ---
    for section_title, bullets in template['static_sections']:
        doc.add_paragraph(section_title, style='EngSubtitle')
        for item in bullets:
            p = doc.add_paragraph(item, style='List Bullet')
            p.style = 'EngBody'
    # --- Table ---
    doc.add_paragraph('Overview Table', style='EngSubtitle')
    table_headers = template['table_headers']
    ai_table_data = ai_sections.get('table_data')
    if not ai_table_data or not isinstance(ai_table_data, list) or not all(isinstance(row, dict) for row in ai_table_data):
        ai_table_data = [dict(zip(table_headers, row)) for row in template['table_data']]
    table = doc.add_table(rows=1, cols=len(table_headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for i, hdr in enumerate(table_headers):
        cell_p = hdr_cells[i].paragraphs[0]
        cell_p.text = hdr
        cell_p.style = 'EngTableHeader'
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'D3D3D3')  # Light gray background
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
        hdr_cells[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for idx, row_data in enumerate(ai_table_data):
        row = table.add_row().cells
        for i, hdr in enumerate(table_headers):
            val = row_data.get(hdr, '')
            # Format as price if header is cost-related
            if 'cost' in hdr.lower() or 'budget' in hdr.lower() or 'price' in hdr.lower():
                row[i].text = format_price(val)
            else:
                row[i].text = str(val)
            row[i].paragraphs[0].style = 'EngBody'
            row[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if idx % 2 == 1:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F2F6FA')
                row[i]._tc.get_or_add_tcPr().append(shading_elm)
    # Add spacing after table
    doc.add_paragraph()
    # --- Branding ---
    doc.add_paragraph('Company Branding', style='EngSubtitle')
    if logo_path:
        try:
            logo_brand_par = doc.add_paragraph()
            logo_brand_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            add_scaled_logo(logo_brand_par, logo_path, max_width=Inches(1.5), max_pixel_height=40)
        except Exception:
            doc.add_paragraph("[Logo could not be displayed]")
    else:
        insert_logos(doc)
    # --- Hyperlink ---
    p = doc.add_paragraph('For more, visit ', style='EngBody')
    add_hyperlink(p, 'VibeOps.ca - ', 'https://www.vibeops.ca')
    # --- Appendix ---
    doc.add_page_break()
    doc.add_paragraph('Appendix: Supporting Data', style='EngSubtitle')
    doc.add_paragraph('Detailed cost breakdowns, risk assessments, and technical specifications are available upon request.', style='EngBody')
    # --- Save ---
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f'vibeops_{report_type}_report_{timestamp}.docx'
    doc.save(filename)
    return filename 
